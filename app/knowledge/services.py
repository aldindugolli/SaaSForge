import json
import os
import re
from datetime import UTC, date, datetime, timedelta
from io import BytesIO
from uuid import uuid4

from flask import current_app
from sqlalchemy import func

from app.core.extensions import cache, db
from app.core.models import NotificationType, Organization, User
from app.knowledge.models import (
    DocumentChunk,
    KnowledgeCollection,
    KnowledgeConversation,
    KnowledgeDocument,
    KnowledgeMessage,
    KnowledgeUsage,
)
from app.services.audit_service import AuditService
from app.services.notification_service import NotificationService


def _org_dir(org_id):
    base = current_app.config.get("KNOWLEDGE_UPLOAD_FOLDER", os.path.join(current_app.instance_path, "knowledge"))
    path = os.path.join(base, org_id)
    os.makedirs(path, exist_ok=True)
    return path


class DocumentProcessor:
    SUPPORTED_TYPES = {"pdf", "docx", "txt", "md"}

    @staticmethod
    def extract_text(file_path: str, file_type: str) -> tuple[str, int, int]:
        if file_type == "pdf":
            return DocumentProcessor._extract_pdf(file_path)
        elif file_type == "docx":
            return DocumentProcessor._extract_docx(file_path)
        else:
            return DocumentProcessor._extract_text(file_path)

    @staticmethod
    def _extract_pdf(file_path: str) -> tuple[str, int, int]:
        try:
            import fitz
        except ImportError:
            raise ImportError("PyMuPDF (fitz) is required for PDF processing")
        doc = fitz.open(file_path)
        text = "\n".join(page.get_text() for page in doc)
        page_count = len(doc)
        doc.close()
        word_count = len(text.split())
        return text, word_count, page_count

    @staticmethod
    def _extract_docx(file_path: str) -> tuple[str, int, int]:
        try:
            from docx import Document as DocxDocument
        except ImportError:
            raise ImportError("python-docx is required for DOCX processing")
        doc = DocxDocument(file_path)
        text = "\n".join(p.text for p in doc.paragraphs)
        word_count = len(text.split())
        return text, word_count, 0

    @staticmethod
    def _extract_text(file_path: str) -> tuple[str, int, int]:
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        word_count = len(text.split())
        return text, word_count, 0

    @staticmethod
    def chunk_text(text: str, chunk_size: int = 1000, overlap: int = 200) -> list[str]:
        words = text.split()
        chunks = []
        start = 0
        while start < len(words):
            end = start + chunk_size
            chunk = " ".join(words[start:end])
            chunks.append(chunk)
            start += chunk_size - overlap
        return chunks or [text]


class AIProvider:
    def generate_embedding(self, text: str) -> list[float]:
        raise NotImplementedError

    def chat_completion(
        self, messages: list[dict], context_chunks: list[str] | None = None
    ) -> tuple[str, int]:
        raise NotImplementedError


class OpenAIProvider(AIProvider):
    def __init__(self):
        import openai
        self.client = openai.OpenAI(api_key=os.environ.get("OPENAI_API_KEY", ""))

    def generate_embedding(self, text: str) -> list[float]:
        resp = self.client.embeddings.create(
            model="text-embedding-3-small", input=text
        )
        return resp.data[0].embedding

    def chat_completion(
        self, messages: list[dict], context_chunks: list[str] | None = None
    ) -> tuple[str, int]:
        if context_chunks:
            context = "\n\n".join(
                f"[Source {i+1}]: {chunk}" for i, chunk in enumerate(context_chunks)
            )
            system_msg = {
                "role": "system",
                "content": (
                    "You are a helpful AI assistant that answers questions based on the provided documents. "
                    "Use the following document excerpts to answer the user's question. "
                    "If the answer cannot be found in the documents, say so clearly. "
                    "Always cite the specific source numbers you used.\n\n"
                    f"DOCUMENT EXCERPTS:\n{context}"
                ),
            }
        else:
            system_msg = {
                "role": "system",
                "content": "You are a helpful AI assistant for a knowledge base platform.",
            }
        all_messages = [system_msg] + messages
        resp = self.client.chat.completions.create(
            model="gpt-4o-mini",
            messages=all_messages,
            max_tokens=2048,
        )
        return resp.choices[0].message.content, resp.usage.total_tokens


class MockProvider(AIProvider):
    def generate_embedding(self, text: str) -> list[float]:
        import hashlib
        h = hashlib.sha256(text.encode()).digest()
        norm = sum(b * b for b in h) ** 0.5
        return [b / norm for b in h[:128]]

    def chat_completion(
        self, messages: list[dict], context_chunks: list[str] | None = None
    ) -> tuple[str, int]:
        user_msg = next((m["content"] for m in messages if m["role"] == "user"), "")
        if context_chunks:
            citations = "\n\n".join(
                f"> Source {i+1}: {chunk[:200]}..."
                for i, chunk in enumerate(context_chunks[:3])
            )
            response = (
                f"Based on the provided documents, here is my response to: '{user_msg[:100]}...'\n\n"
                f"I found relevant information across {len(context_chunks)} document sections. "
                f"The key points from the documents suggest that this is a well-documented topic.\n\n"
                f"{citations}\n\n"
                f"*This is a demo response. Set OPENAI_API_KEY to enable AI-powered responses.*"
            )
        else:
            response = (
                f"Hello! I'm your knowledge base assistant. "
                f"You asked: '{user_msg[:100]}...'\n\n"
                f"This is a demo response. Upload documents and set OPENAI_API_KEY for real AI responses."
            )
        return response, 50 + len(user_msg)


def _get_ai_provider() -> AIProvider:
    provider = os.environ.get("AI_PROVIDER", "mock")
    if provider == "openai" and os.environ.get("OPENAI_API_KEY"):
        return OpenAIProvider()
    return MockProvider()


class KnowledgeService:
    _cache_prefix = "knowledge"

    @staticmethod
    def _cache_key(name: str, org_id: str, **params) -> str:
        parts = [f"org={org_id}", name]
        for k, v in sorted(params.items()):
            parts.append(f"{k}={v}")
        return ":".join(parts)

    @staticmethod
    def _track_usage(org_id: str, field: str, amount: int = 1):
        today = date.today()
        usage = KnowledgeUsage.query.filter_by(
            organization_id=org_id, date=today
        ).first()
        if not usage:
            usage = KnowledgeUsage(organization_id=org_id, date=today)
            db.session.add(usage)
        setattr(usage, field, getattr(usage, field, 0) + amount)
        db.session.commit()

    @staticmethod
    def _invalidate_org_cache(org_id: str):
        cache.invalidate_pattern(f"{KnowledgeService._cache_prefix}:org={org_id}:*")

    @staticmethod
    def get_usage_stats(org_id: str) -> dict:
        today = date.today()
        thirty_days_ago = today - timedelta(days=30)

        current = KnowledgeUsage.query.filter_by(
            organization_id=org_id, date=today
        ).first()

        monthly = db.session.query(
            func.sum(KnowledgeUsage.documents_uploaded).label("documents"),
            func.sum(KnowledgeUsage.searches_performed).label("searches"),
            func.sum(KnowledgeUsage.chat_messages).label("messages"),
            func.sum(KnowledgeUsage.ai_tokens_used).label("tokens"),
        ).filter(
            KnowledgeUsage.organization_id == org_id,
            KnowledgeUsage.date >= thirty_days_ago,
        ).first()

        doc_count = KnowledgeDocument.query.filter_by(
            organization_id=org_id
        ).count()
        convo_count = KnowledgeConversation.query.filter_by(
            organization_id=org_id
        ).count()

        return {
            "today": {
                "documents_uploaded": current.documents_uploaded if current else 0,
                "searches_performed": current.searches_performed if current else 0,
                "chat_messages": current.chat_messages if current else 0,
                "ai_tokens_used": current.ai_tokens_used if current else 0,
            },
            "monthly": {
                "documents": monthly.documents or 0,
                "searches": monthly.searches or 0,
                "messages": monthly.messages or 0,
                "tokens": monthly.tokens or 0,
            },
            "totals": {
                "documents": doc_count,
                "conversations": convo_count,
            },
        }

    @staticmethod
    def get_usage_analytics(org_id: str, days: int = 30) -> dict:
        start_date = date.today() - timedelta(days=days)
        records = (
            KnowledgeUsage.query.filter(
                KnowledgeUsage.organization_id == org_id,
                KnowledgeUsage.date >= start_date,
            )
            .order_by(KnowledgeUsage.date)
            .all()
        )
        dates = []
        documents = []
        searches = []
        messages = []
        tokens = []
        for r in records:
            dates.append(r.date.isoformat())
            documents.append(r.documents_uploaded)
            searches.append(r.searches_performed)
            messages.append(r.chat_messages)
            tokens.append(r.ai_tokens_used)
        return {
            "dates": dates,
            "documents": documents,
            "searches": searches,
            "messages": messages,
            "tokens": tokens,
        }

    @staticmethod
    def get_dashboard_stats(org_id: str) -> dict:
        ck = KnowledgeService._cache_key("dashboard", org_id)
        cached = cache.get(KnowledgeService._cache_prefix, ck)
        if cached is not None:
            return cached

        total_docs = KnowledgeDocument.query.filter_by(organization_id=org_id).count()
        total_collections = KnowledgeCollection.query.filter_by(
            organization_id=org_id
        ).count()
        total_conversations = KnowledgeConversation.query.filter_by(
            organization_id=org_id
        ).count()
        total_chunks = (
            db.session.query(func.count(DocumentChunk.id))
            .join(KnowledgeDocument)
            .filter(KnowledgeDocument.organization_id == org_id)
            .scalar()
        ) or 0
        recent_docs = (
            KnowledgeDocument.query.filter_by(organization_id=org_id)
            .order_by(KnowledgeDocument.created_at.desc())
            .limit(5)
            .all()
        )
        recent_convs = (
            KnowledgeConversation.query.filter_by(organization_id=org_id)
            .order_by(KnowledgeConversation.updated_at.desc())
            .limit(5)
            .all()
        )

        docs_by_type = (
            db.session.query(
                KnowledgeDocument.type,
                func.count(KnowledgeDocument.id).label("count"),
            )
            .filter(KnowledgeDocument.organization_id == org_id)
            .group_by(KnowledgeDocument.type)
            .all()
        )

        by_collection = (
            db.session.query(
                KnowledgeCollection.name,
                func.count(KnowledgeDocument.id).label("count"),
            )
            .join(KnowledgeDocument, KnowledgeDocument.collection_id == KnowledgeCollection.id, isouter=True)
            .filter(KnowledgeCollection.organization_id == org_id)
            .group_by(KnowledgeCollection.name)
            .all()
        )

        result = {
            "total_documents": total_docs,
            "total_collections": total_collections,
            "total_conversations": total_conversations,
            "total_chunks": total_chunks,
            "recent_documents": [
                {
                    "id": d.id,
                    "name": d.name,
                    "type": d.type,
                    "size": d.size,
                    "word_count": d.word_count,
                    "created_at": d.created_at.isoformat(),
                }
                for d in recent_docs
            ],
            "recent_conversations": [
                {
                    "id": c.id,
                    "title": c.title,
                    "message_count": c.message_count,
                    "created_at": c.created_at.isoformat(),
                }
                for c in recent_convs
            ],
            "documents_by_type": {r.type: r.count for r in docs_by_type},
            "documents_by_collection": {r.name: r.count for r in by_collection},
        }
        cache.set(KnowledgeService._cache_prefix, ck, result, 300)
        return result

    @staticmethod
    def create_collection(
        org_id: str, user_id: str, name: str, description: str | None = None
    ) -> KnowledgeCollection:
        if KnowledgeCollection.query.filter_by(
            organization_id=org_id, name=name
        ).first():
            raise ValueError("A collection with this name already exists")
        collection = KnowledgeCollection(
            organization_id=org_id,
            name=name,
            description=description,
            created_by_id=user_id,
        )
        db.session.add(collection)
        db.session.commit()
        KnowledgeService._invalidate_org_cache(org_id)
        AuditService.log(
            actor_id=user_id,
            organization_id=org_id,
            action="knowledge.collection.create",
            resource_type="knowledge_collection",
            resource_id=collection.id,
            metadata={"name": name},
        )
        return collection

    @staticmethod
    def list_collections(org_id: str) -> list[KnowledgeCollection]:
        return (
            KnowledgeCollection.query.filter_by(organization_id=org_id)
            .order_by(KnowledgeCollection.name)
            .all()
        )

    @staticmethod
    def delete_collection(org_id: str, collection_id: str, user_id: str):
        collection = KnowledgeCollection.query.filter_by(
            id=collection_id, organization_id=org_id
        ).first_or_404()
        db.session.delete(collection)
        db.session.commit()
        KnowledgeService._invalidate_org_cache(org_id)
        AuditService.log(
            actor_id=user_id,
            organization_id=org_id,
            action="knowledge.collection.delete",
            resource_type="knowledge_collection",
            resource_id=collection_id,
        )

    @staticmethod
    def upload_document(
        org_id: str,
        user_id: str,
        name: str,
        file_data: bytes,
        file_type: str,
        collection_id: str | None = None,
        tags: list[str] | None = None,
    ) -> KnowledgeDocument:
        if file_type not in DocumentProcessor.SUPPORTED_TYPES:
            raise ValueError(f"Unsupported file type: {file_type}")

        org_dir = _org_dir(org_id)
        file_id = str(uuid4())
        safe_name = re.sub(r"[^\w\-_. ]", "_", name)
        file_path = os.path.join(org_dir, f"{file_id}_{safe_name}")
        with open(file_path, "wb") as f:
            f.write(file_data)

        doc = KnowledgeDocument(
            organization_id=org_id,
            collection_id=collection_id,
            name=name,
            type=file_type,
            size=len(file_data),
            file_path=file_path,
            tags=tags or [],
            uploaded_by_id=user_id,
            embedding_status="pending",
        )
        db.session.add(doc)
        db.session.commit()

        if collection_id:
            collection = KnowledgeCollection.query.get(collection_id)
            if collection:
                collection.document_count = (
                    KnowledgeDocument.query.filter_by(collection_id=collection_id).count()
                )
                db.session.commit()

        KnowledgeService._invalidate_org_cache(org_id)
        KnowledgeService._track_usage(org_id, "documents_uploaded")

        AuditService.log(
            actor_id=user_id,
            organization_id=org_id,
            action="knowledge.document.upload",
            resource_type="knowledge_document",
            resource_id=doc.id,
            metadata={"name": name, "type": file_type, "size": len(file_data)},
        )

        org = Organization.query.get(org_id)
        members = org.members if org else []
        notifications = [
            {
                "user_id": m.id,
                "title": f"New document: {name}",
                "message": f"{name} was uploaded to your knowledge base",
                "type": NotificationType.INFO.value,
                "link": f"/knowledge/documents/{doc.id}",
            }
            for m in members
            if m.id != user_id
        ]
        if notifications:
            NotificationService.bulk_create(notifications)

        return doc

    @staticmethod
    def process_document(doc_id: str) -> dict:
        doc = KnowledgeDocument.query.get(doc_id)
        if not doc:
            raise ValueError("Document not found")
        if not doc.file_path or not os.path.exists(doc.file_path):
            raise ValueError("Document file not found")

        doc.embedding_status = "processing"
        db.session.commit()

        try:
            text, word_count, page_count = DocumentProcessor.extract_text(
                doc.file_path, doc.type
            )
            chunks = DocumentProcessor.chunk_text(text)

            DocumentChunk.query.filter_by(document_id=doc.id).delete()
            ai = _get_ai_provider()

            for i, chunk_text in enumerate(chunks):
                embedding = ai.generate_embedding(chunk_text)
                chunk = DocumentChunk(
                    document_id=doc.id,
                    chunk_index=i,
                    content=chunk_text,
                    embedding=embedding,
                    token_count=len(chunk_text.split()),
                )
                db.session.add(chunk)

            doc.content_text = text
            doc.word_count = word_count
            doc.page_count = page_count
            doc.chunk_count = len(chunks)
            doc.embedding_status = "completed"
            db.session.commit()
            KnowledgeService._invalidate_org_cache(doc.organization_id)

            return {
                "status": "completed",
                "chunks": len(chunks),
                "words": word_count,
                "pages": page_count,
            }
        except Exception as e:
            doc.embedding_status = "failed"
            db.session.commit()
            raise e

    @staticmethod
    def get_document(org_id: str, doc_id: str) -> KnowledgeDocument | None:
        return KnowledgeDocument.query.filter_by(
            id=doc_id, organization_id=org_id
        ).first()

    @staticmethod
    def list_documents(
        org_id: str,
        collection_id: str | None = None,
        search: str | None = None,
        tag: str | None = None,
        type_filter: str | None = None,
        page: int = 1,
        per_page: int = 20,
    ) -> tuple[list[KnowledgeDocument], int, int]:
        query = KnowledgeDocument.query.filter_by(organization_id=org_id)
        if collection_id:
            query = query.filter_by(collection_id=collection_id)
        if type_filter:
            query = query.filter_by(type=type_filter)
        if tag:
            query = query.filter(KnowledgeDocument.tags.contains(tag))
        if search:
            like = f"%{search}%"
            query = query.filter(
                db.or_(
                    KnowledgeDocument.name.ilike(like),
                    KnowledgeDocument.content_text.ilike(like),
                )
            )
        total = query.count()
        total_pages = max(1, (total + per_page - 1) // per_page)
        documents = (
            query.order_by(KnowledgeDocument.created_at.desc())
            .offset((page - 1) * per_page)
            .limit(per_page)
            .all()
        )
        return documents, total, total_pages

    @staticmethod
    def delete_document(org_id: str, doc_id: str, user_id: str):
        doc = KnowledgeDocument.query.filter_by(
            id=doc_id, organization_id=org_id
        ).first_or_404()
        if doc.file_path and os.path.exists(doc.file_path):
            os.remove(doc.file_path)
        db.session.delete(doc)
        db.session.commit()
        KnowledgeService._invalidate_org_cache(org_id)
        AuditService.log(
            actor_id=user_id,
            organization_id=org_id,
            action="knowledge.document.delete",
            resource_type="knowledge_document",
            resource_id=doc_id,
            metadata={"name": doc.name},
        )

    @staticmethod
    def update_document(
        org_id: str,
        doc_id: str,
        user_id: str,
        name: str | None = None,
        collection_id: str | None = None,
        tags: list[str] | None = None,
    ) -> KnowledgeDocument:
        doc = KnowledgeDocument.query.filter_by(
            id=doc_id, organization_id=org_id
        ).first_or_404()
        if name:
            doc.name = name
        if collection_id is not None:
            doc.collection_id = collection_id or None
        if tags is not None:
            doc.tags = tags
        db.session.commit()
        KnowledgeService._invalidate_org_cache(org_id)
        AuditService.log(
            actor_id=user_id,
            organization_id=org_id,
            action="knowledge.document.update",
            resource_type="knowledge_document",
            resource_id=doc_id,
        )
        return doc

    @staticmethod
    def search(
        org_id: str,
        query: str,
        collection_id: str | None = None,
        use_semantic: bool = True,
        limit: int = 10,
    ) -> list[dict]:
        KnowledgeService._track_usage(org_id, "searches_performed")

        base_filter = DocumentChunk.query.join(KnowledgeDocument).filter(
            KnowledgeDocument.organization_id == org_id
        )
        if collection_id:
            base_filter = base_filter.filter(
                KnowledgeDocument.collection_id == collection_id
            )

        if use_semantic:
            try:
                ai = _get_ai_provider()
                query_embedding = ai.generate_embedding(query)
                all_chunks = base_filter.filter(
                    DocumentChunk.embedding.isnot(None)
                ).all()

                def cosine_sim(a, b):
                    dot = sum(x * y for x, y in zip(a, b))
                    na = sum(x * x for x in a) ** 0.5
                    nb = sum(y * y for y in b) ** 0.5
                    return dot / (na * nb) if na and nb else 0

                scored = []
                for chunk in all_chunks:
                    if chunk.embedding:
                        sim = cosine_sim(query_embedding, chunk.embedding)
                        scored.append((sim, chunk))
                scored.sort(key=lambda x: x[0], reverse=True)
                top_chunks = scored[:limit]
            except Exception:
                use_semantic = False

        if not use_semantic:
            like = f"%{query}%"
            top_chunks = [
                (0.0, c)
                for c in base_filter.filter(DocumentChunk.content.ilike(like))
                .order_by(DocumentChunk.chunk_index)
                .limit(limit)
                .all()
            ]

        results = []
        seen_docs = set()
        for score, chunk in top_chunks:
            doc = KnowledgeDocument.query.get(chunk.document_id)
            if doc and doc.id not in seen_docs:
                seen_docs.add(doc.id)
                doc.search_count = (doc.search_count or 0) + 1
                results.append(
                    {
                        "document_id": doc.id,
                        "document_name": doc.name,
                        "document_type": doc.type,
                        "chunk_index": chunk.chunk_index,
                        "content": chunk.content[:500],
                        "score": round(score, 4) if use_semantic else None,
                        "document_url": f"/knowledge/documents/{doc.id}",
                    }
                )
        db.session.commit()
        return results

    @staticmethod
    def create_conversation(
        org_id: str,
        user_id: str,
        title: str,
        document_id: str | None = None,
    ) -> KnowledgeConversation:
        conversation = KnowledgeConversation(
            organization_id=org_id,
            document_id=document_id,
            title=title,
            created_by_id=user_id,
        )
        db.session.add(conversation)
        db.session.commit()
        KnowledgeService._invalidate_org_cache(org_id)
        return conversation

    @staticmethod
    def get_conversation(org_id: str, conversation_id: str):
        return KnowledgeConversation.query.filter_by(
            id=conversation_id, organization_id=org_id
        ).first()

    @staticmethod
    def list_conversations(
        org_id: str, page: int = 1, per_page: int = 20
    ) -> tuple[list[KnowledgeConversation], int, int]:
        query = KnowledgeConversation.query.filter_by(organization_id=org_id)
        total = query.count()
        total_pages = max(1, (total + per_page - 1) // per_page)
        conversations = (
            query.order_by(KnowledgeConversation.updated_at.desc())
            .offset((page - 1) * per_page)
            .limit(per_page)
            .all()
        )
        return conversations, total, total_pages

    @staticmethod
    def send_message(
        org_id: str,
        conversation_id: str,
        user_id: str,
        content: str,
    ) -> tuple[KnowledgeMessage, KnowledgeMessage]:
        conversation = KnowledgeConversation.query.filter_by(
            id=conversation_id, organization_id=org_id
        ).first_or_404()

        user_msg = KnowledgeMessage(
            conversation_id=conversation_id, role="user", content=content
        )
        db.session.add(user_msg)

        search_results = KnowledgeService.search(
            org_id, content, limit=5, use_semantic=True
        )
        context_chunks = [r["content"] for r in search_results]

        ai = _get_ai_provider()
        history = (
            KnowledgeMessage.query.filter_by(conversation_id=conversation_id)
            .order_by(KnowledgeMessage.created_at)
            .all()
        )
        messages = [
            {"role": m.role, "content": m.content} for m in history[-10:]
        ]
        messages.append({"role": "user", "content": content})

        response_text, tokens_used = ai.chat_completion(messages, context_chunks)

        citations = [
            {
                "document_id": r["document_id"],
                "document_name": r["document_name"],
                "snippet": r["content"][:200],
            }
            for r in search_results[:3]
        ]

        assistant_msg = KnowledgeMessage(
            conversation_id=conversation_id,
            role="assistant",
            content=response_text,
            citations=citations,
            token_count=tokens_used,
        )
        db.session.add(assistant_msg)

        conversation.message_count = (
            KnowledgeMessage.query.filter_by(conversation_id=conversation_id).count()
        )
        conversation.updated_at = datetime.now(UTC)
        db.session.commit()

        KnowledgeService._track_usage(org_id, "chat_messages")
        KnowledgeService._track_usage(org_id, "ai_tokens_used", tokens_used)

        return user_msg, assistant_msg

    @staticmethod
    def get_conversation_messages(conversation_id: str) -> list[KnowledgeMessage]:
        return (
            KnowledgeMessage.query.filter_by(conversation_id=conversation_id)
            .order_by(KnowledgeMessage.created_at)
            .all()
        )

    @staticmethod
    def delete_conversation(org_id: str, conversation_id: str, user_id: str):
        conversation = KnowledgeConversation.query.filter_by(
            id=conversation_id, organization_id=org_id
        ).first_or_404()
        db.session.delete(conversation)
        db.session.commit()
        KnowledgeService._invalidate_org_cache(org_id)
